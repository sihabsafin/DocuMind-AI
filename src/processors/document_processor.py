"""
Document processors for various file formats
Handles PDF, DOCX, TXT, and Markdown files
"""
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from abc import ABC, abstractmethod

import PyPDF2
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup


@dataclass
class DocumentSection:
    """Represents a section in a document"""
    title: str
    content: str
    level: int  # Heading level (1-6)
    page_number: Optional[int] = None
    start_position: Optional[int] = None
    end_position: Optional[int] = None


@dataclass
class ProcessedDocument:
    """Structured representation of a processed document"""
    filename: str
    file_type: str
    full_text: str
    sections: List[DocumentSection]
    metadata: Dict
    tables: List[Dict]
    total_pages: Optional[int] = None
    total_chars: int = 0
    total_words: int = 0


class BaseDocumentProcessor(ABC):
    """Base class for document processors"""
    
    @abstractmethod
    def process(self, file_path: Path) -> ProcessedDocument:
        """Process a document and return structured data"""
        pass
    
    def extract_sections(self, text: str) -> List[DocumentSection]:
        """Extract sections from text based on headings"""
        sections = []
        
        # Regex patterns for different heading formats
        patterns = [
            (r'^#{1,6}\s+(.+)$', 'markdown'),  # Markdown headers
            (r'^([A-Z][A-Z\s]{3,})\s*$', 'all_caps'),  # ALL CAPS headers
            (r'^(\d+\.(?:\d+\.)*)\s+(.+)$', 'numbered'),  # 1. 1.1. headers
            (r'^([IVXLCDM]+\.\s+.+)$', 'roman'),  # Roman numerals
        ]
        
        lines = text.split('\n')
        current_section = None
        content_buffer = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            is_heading = False
            heading_level = 3  # Default level
            heading_text = line_stripped
            
            # Check against patterns
            for pattern, pattern_type in patterns:
                match = re.match(pattern, line_stripped, re.MULTILINE)
                if match:
                    is_heading = True
                    if pattern_type == 'markdown':
                        heading_level = len(re.match(r'^#+', line_stripped).group())
                        heading_text = match.group().lstrip('#').strip()
                    elif pattern_type == 'numbered':
                        heading_level = len(match.group(1).split('.'))
                        heading_text = match.group(2)
                    elif pattern_type in ['all_caps', 'roman']:
                        heading_level = 2
                        heading_text = match.group(1) if pattern_type == 'roman' else match.group()
                    break
            
            if is_heading:
                # Save previous section
                if current_section:
                    current_section.content = '\n'.join(content_buffer).strip()
                    sections.append(current_section)
                    content_buffer = []
                
                # Start new section
                current_section = DocumentSection(
                    title=heading_text,
                    content="",
                    level=heading_level,
                    start_position=i
                )
            else:
                content_buffer.append(line)
        
        # Save last section
        if current_section:
            current_section.content = '\n'.join(content_buffer).strip()
            sections.append(current_section)
        
        # If no sections found, create a default one
        if not sections:
            sections.append(DocumentSection(
                title="Document Content",
                content=text,
                level=1
            ))
        
        return sections
    
    def count_words(self, text: str) -> int:
        """Count words in text"""
        return len(re.findall(r'\b\w+\b', text))


class PDFProcessor(BaseDocumentProcessor):
    """Process PDF files"""
    
    def process(self, file_path: Path) -> ProcessedDocument:
        """Process a PDF file"""
        full_text = ""
        tables = []
        total_pages = 0
        
        try:
            # Try pdfplumber first (better for tables and layout)
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n\n--- Page {page_num} ---\n\n"
                        full_text += page_text
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        tables.append({
                            'page': page_num,
                            'data': table,
                            'description': self._describe_table(table)
                        })
        
        except Exception as e:
            # Fallback to PyPDF2
            print(f"pdfplumber failed, using PyPDF2: {e}")
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                total_pages = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n\n--- Page {page_num} ---\n\n"
                        full_text += page_text
        
        # Extract sections
        sections = self.extract_sections(full_text)
        
        # Create metadata
        metadata = {
            'total_pages': total_pages,
            'has_tables': len(tables) > 0,
            'table_count': len(tables)
        }
        
        return ProcessedDocument(
            filename=file_path.name,
            file_type='pdf',
            full_text=full_text,
            sections=sections,
            metadata=metadata,
            tables=tables,
            total_pages=total_pages,
            total_chars=len(full_text),
            total_words=self.count_words(full_text)
        )
    
    def _describe_table(self, table: List[List]) -> str:
        """Generate a text description of a table"""
        if not table or not table[0]:
            return ""
        
        # Assume first row is header
        header = table[0]
        num_rows = len(table) - 1
        num_cols = len(header)
        
        description = f"Table with {num_rows} rows and {num_cols} columns. "
        description += f"Columns: {', '.join(str(h) for h in header if h)}. "
        
        return description


class DOCXProcessor(BaseDocumentProcessor):
    """Process DOCX files"""
    
    def process(self, file_path: Path) -> ProcessedDocument:
        """Process a DOCX file"""
        doc = Document(file_path)
        
        full_text = ""
        sections = []
        tables = []
        
        current_section = None
        content_buffer = []
        
        for element in doc.element.body:
            # Process paragraphs
            if element.tag.endswith('p'):
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.replace('Heading ', ''))
                        
                        # Save previous section
                        if current_section:
                            current_section.content = '\n'.join(content_buffer).strip()
                            sections.append(current_section)
                            content_buffer = []
                        
                        # Start new section
                        current_section = DocumentSection(
                            title=text,
                            content="",
                            level=level
                        )
                    else:
                        content_buffer.append(text)
                        full_text += text + "\n"
            
            # Process tables
            elif element.tag.endswith('tbl'):
                table = next((t for t in doc.tables if t._element == element), None)
                if table:
                    table_data = [[cell.text for cell in row.cells] for row in table.rows]
                    tables.append({
                        'data': table_data,
                        'description': self._describe_table(table_data)
                    })
                    
                    # Add table description to text
                    table_text = f"\n[TABLE: {self._describe_table(table_data)}]\n"
                    content_buffer.append(table_text)
                    full_text += table_text
        
        # Save last section
        if current_section:
            current_section.content = '\n'.join(content_buffer).strip()
            sections.append(current_section)
        
        # If no sections, create default
        if not sections:
            sections.append(DocumentSection(
                title="Document Content",
                content=full_text,
                level=1
            ))
        
        metadata = {
            'has_tables': len(tables) > 0,
            'table_count': len(tables),
            'paragraph_count': len(doc.paragraphs)
        }
        
        return ProcessedDocument(
            filename=file_path.name,
            file_type='docx',
            full_text=full_text,
            sections=sections,
            metadata=metadata,
            tables=tables,
            total_chars=len(full_text),
            total_words=self.count_words(full_text)
        )
    
    def _describe_table(self, table: List[List]) -> str:
        """Generate a text description of a table"""
        if not table or not table[0]:
            return ""
        
        header = table[0]
        num_rows = len(table) - 1
        num_cols = len(header)
        
        description = f"Table with {num_rows} rows and {num_cols} columns. "
        description += f"Columns: {', '.join(str(h) for h in header if h)}. "
        
        return description


class TextProcessor(BaseDocumentProcessor):
    """Process plain text and markdown files"""
    
    def process(self, file_path: Path) -> ProcessedDocument:
        """Process a text or markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            full_text = f.read()
        
        sections = self.extract_sections(full_text)
        
        metadata = {
            'encoding': 'utf-8',
            'is_markdown': file_path.suffix.lower() in ['.md', '.markdown']
        }
        
        return ProcessedDocument(
            filename=file_path.name,
            file_type=file_path.suffix.lower().lstrip('.'),
            full_text=full_text,
            sections=sections,
            metadata=metadata,
            tables=[],
            total_chars=len(full_text),
            total_words=self.count_words(full_text)
        )


class DocumentProcessorFactory:
    """Factory to get the appropriate processor for a file"""
    
    @staticmethod
    def get_processor(file_path: Path) -> BaseDocumentProcessor:
        """Get the appropriate processor based on file extension"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return PDFProcessor()
        elif suffix == '.docx':
            return DOCXProcessor()
        elif suffix in ['.txt', '.md', '.markdown']:
            return TextProcessor()
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    @staticmethod
    def process_document(file_path: Path) -> ProcessedDocument:
        """Process a document with the appropriate processor"""
        processor = DocumentProcessorFactory.get_processor(file_path)
        return processor.process(file_path)
